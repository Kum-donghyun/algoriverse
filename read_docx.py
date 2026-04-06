from docx import Document

doc = Document(r'C:\Users\smhrd\Desktop\back\01. 실시간 이슈 추적 및 요약 서비스.docx')

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] [{para.style.name}] {para.text}")

print("\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i} ---")
    for row in table.rows:
        cells = [cell.text.strip()[:100] for cell in row.cells]
        print(" | ".join(cells))
